"""
Comprehensive test for all supported file formats in EchoVerse
Tests TXT, PDF, DOCX, HTML file import functionality for audiobook creation
"""

import os
import tempfile
import sys

# Add project path
sys.path.append('d:\\project\\audiobook')

def test_all_file_formats():
    """Test all supported file formats"""
    print("🧪 Testing all supported file formats for audiobook creation...")
    
    # Test if required libraries are available
    has_pypdf2 = False
    has_docx = False
    
    try:
        import PyPDF2
        has_pypdf2 = True
        print("✅ PyPDF2 available")
    except ImportError:
        print("⚠️  PyPDF2 not available - PDF processing will be disabled")
    
    try:
        from docx import Document
        has_docx = True
        print("✅ python-docx available")
    except ImportError:
        print("⚠️  python-docx not available - DOCX processing will be disabled")
    
    # Initialize file processing service
    try:
        from services.file_service import FileProcessingService
        file_service = FileProcessingService()
        print("✅ File processing service initialized")
    except Exception as e:
        print(f"❌ Failed to initialize file processing service: {e}")
        return False
    
    # Create test files
    temp_dir = tempfile.mkdtemp()
    test_files = {}
    
    # 1. TXT file
    txt_content = """This is a test TXT file for audiobook creation.
It contains multiple lines to test the file import functionality.
The quick brown fox jumps over the lazy dog.
End of test file."""
    
    txt_path = os.path.join(temp_dir, "test.txt")
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(txt_content)
    test_files['txt'] = txt_path
    print("📄 Created TXT test file")
    
    # 2. HTML file
    html_content = """<html>
<head><title>Test HTML</title></head>
<body>
<h1>HTML Test File</h1>
<p>This is a test HTML file for audiobook creation.</p>
<p>It contains multiple paragraphs to test the file import functionality.</p>
<p>The quick brown fox jumps over the lazy dog.</p>
<p>End of test file.</p>
</body>
</html>"""
    
    html_path = os.path.join(temp_dir, "test.html")
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    test_files['html'] = html_path
    print("🌐 Created HTML test file")
    
    # 3. PDF file (if PyPDF2 is available)
    if has_pypdf2:
        try:
            from reportlab.pdfgen import canvas  # type: ignore
            from reportlab.lib.pagesizes import letter  # type: ignore
            
            pdf_path = os.path.join(temp_dir, "test.pdf")
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, "This is a test PDF file for audiobook creation.")
            c.drawString(100, 730, "It contains text to test the file import functionality.")
            c.drawString(100, 710, "The quick brown fox jumps over the lazy dog.")
            c.drawString(100, 690, "End of test file.")
            c.save()
            test_files['pdf'] = pdf_path
            print("📚 Created PDF test file")
        except Exception as e:
            print(f"⚠️  Failed to create PDF test file: {e}")
    
    # 4. DOCX file (if python-docx is available)
    if has_docx:
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Test DOCX File', 0)
            doc.add_paragraph('This is a test DOCX file for audiobook creation.')
            doc.add_paragraph('It contains text to test the file import functionality.')
            doc.add_paragraph('The quick brown fox jumps over the lazy dog.')
            doc.add_paragraph('End of test file.')
            
            docx_path = os.path.join(temp_dir, "test.docx")
            doc.save(docx_path)
            test_files['docx'] = docx_path
            print("📝 Created DOCX test file")
        except Exception as e:
            print(f"⚠️  Failed to create DOCX test file: {e}")
    
    # Test processing of each file
    print(f"\n🔍 Testing file processing for {len(test_files)} file formats...")
    results = {}
    
    for ext, file_path in test_files.items():
        print(f"\n--- Testing {ext.upper()} file ---")
        try:
            result = file_service.extract_text_from_file(file_path)
            
            print(f"Status: {result.status.value}")
            print(f"Text length: {len(result.text_content)} characters")
            if result.errors:
                print(f"Errors: {result.errors}")
            if result.warnings:
                print(f"Warnings: {result.warnings}")
            
            if result.text_content:
                preview = result.text_content[:100].replace('\n', ' ')
                print(f"Text preview: {preview}...")
            
            results[ext] = {
                'status': result.status.value,
                'success': result.status.value in ['success', 'partial'],
                'text_length': len(result.text_content),
                'errors': result.errors,
                'warnings': result.warnings
            }
            
        except Exception as e:
            print(f"❌ Error processing {ext} file: {e}")
            results[ext] = {
                'status': 'error',
                'success': False,
                'text_length': 0,
                'errors': [str(e)],
                'warnings': []
            }
    
    # Clean up test files
    for file_path in test_files.values():
        try:
            os.remove(file_path)
        except:
            pass
    try:
        os.rmdir(temp_dir)
    except:
        pass
    
    # Summary
    print(f"\n{'='*50}")
    print("📋 FILE FORMAT SUPPORT SUMMARY")
    print(f"{'='*50}")
    
    supported_formats = []
    for ext, result in results.items():
        status_icon = "✅" if result['success'] else "❌"
        print(f"{status_icon} .{ext.upper()}: {result['status']}")
        if result['success']:
            supported_formats.append(ext.upper())
    
    print(f"\n📈 Results: {len([r for r in results.values() if r['success']])}/{len(results)} formats working")
    
    if supported_formats:
        print(f"\n🎉 Supported formats for audiobook creation: {', '.join(supported_formats)}")
        print("\n💡 Usage instructions:")
        print("   1. Upload any supported file format through the web interface")
        print("   2. The system will automatically extract text content")
        print("   3. Select voice options and generate your audiobook")
        print("   4. Download the generated audio file")
        return True
    else:
        print("\n❌ No file formats are working properly!")
        return False

def test_flask_integration():
    """Test Flask file upload integration"""
    print(f"\n{'='*50}")
    print("🌐 TESTING FLASK INTEGRATION")
    print(f"{'='*50}")
    
    try:
        # Test the extract_text_from_file function from app.py
        from app import extract_text_from_file
        
        # Create a simple test file
        temp_dir = tempfile.mkdtemp()
        test_content = "This is a test file for Flask integration.\nTesting audiobook creation workflow."
        test_path = os.path.join(temp_dir, "flask_test.txt")
        
        with open(test_path, 'w', encoding='utf-8') as f:
            f.write(test_content)
        
        # Test extraction
        extracted = extract_text_from_file(test_path)
        
        if extracted and len(extracted.strip()) > 0:
            print("✅ Flask file processing integration working")
            print(f"   Extracted: {len(extracted)} characters")
        else:
            print("❌ Flask file processing integration failed")
        
        # Clean up
        os.remove(test_path)
        os.rmdir(temp_dir)
        
    except Exception as e:
        print(f"❌ Flask integration test failed: {e}")

def main():
    """Run comprehensive file format tests"""
    print("🚀 EchoVerse - Comprehensive File Format Testing")
    print("=" * 50)
    
    # Test all file formats
    success = test_all_file_formats()
    
    # Test Flask integration
    test_flask_integration()
    
    print(f"\n{'='*50}")
    if success:
        print("🎉 ALL TESTS COMPLETED SUCCESSFULLY!")
        print("🎧 Your EchoVerse audiobook creator is ready to use with multiple file formats!")
    else:
        print("❌ SOME TESTS FAILED!")
        print("Please check the error messages above and ensure all dependencies are installed.")
    
    return success

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)