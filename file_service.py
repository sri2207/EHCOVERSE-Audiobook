"""
File Service for document processing and text extraction
"""
import os
import tempfile
import logging
from typing import Dict, List, Optional, Tuple, Any, Union, TYPE_CHECKING
from dataclasses import dataclass
from enum import Enum
import mimetypes

# Document processing imports with proper typing
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    PyPDF2 = None  # type: ignore
    HAS_PYPDF2 = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    DocxDocument = None  # type: ignore
    HAS_DOCX = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    chardet = None  # type: ignore
    HAS_CHARDET = False

logger = logging.getLogger(__name__)

class FileType(Enum):
    """Supported file types"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"
    UNKNOWN = "unknown"

class ProcessingStatus(Enum):
    """File processing status"""
    SUCCESS = "success"
    ERROR = "error"
    PARTIAL = "partial"
    UNSUPPORTED = "unsupported"

@dataclass
class FileMetadata:
    """File metadata information"""
    filename: str
    file_type: FileType
    size_bytes: int
    mime_type: str
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    estimated_reading_time: Optional[int] = None  # in minutes

@dataclass
class ProcessingResult:
    """Result of file processing"""
    status: ProcessingStatus
    text_content: str
    metadata: FileMetadata
    errors: List[str]
    warnings: List[str]
    processing_time: float

class FileProcessingService:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_types = self._get_supported_types()
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        
    def _get_supported_types(self) -> Dict[FileType, bool]:
        """Get supported file types based on available libraries"""
        return {
            FileType.PDF: HAS_PYPDF2,
            FileType.DOCX: HAS_DOCX,
            FileType.TXT: True,
            FileType.HTML: True,
            FileType.RTF: False,  # Would need additional library
            FileType.DOC: False,  # Would need additional library
        }
    
    def detect_file_type(self, filepath: str) -> Tuple[FileType, str]:
        """Detect file type and MIME type"""
        try:
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(filepath)
            if not mime_type:
                mime_type = "application/octet-stream"
            
            # Get file extension
            _, ext = os.path.splitext(filepath.lower())
            ext = ext.lstrip('.')
            
            # Map extensions to file types
            extension_map = {
                'pdf': FileType.PDF,
                'docx': FileType.DOCX,
                'doc': FileType.DOC,
                'txt': FileType.TXT,
                'text': FileType.TXT,
                'rtf': FileType.RTF,
                'html': FileType.HTML,
                'htm': FileType.HTML,
            }
            
            file_type = extension_map.get(ext, FileType.UNKNOWN)
            
            return file_type, mime_type
            
        except Exception as e:
            logger.warning(f"⚠️ File type detection failed: {e}")
            return FileType.UNKNOWN, "application/octet-stream"
    
    def validate_file(self, filepath: str) -> Tuple[bool, List[str]]:
        """Validate file for processing"""
        errors = []
        
        # Check if file exists
        if not os.path.exists(filepath):
            errors.append("File does not exist")
            return False, errors
        
        # Check file size
        try:
            file_size = os.path.getsize(filepath)
            if file_size > self.max_file_size:
                errors.append(f"File too large: {file_size / (1024*1024):.1f}MB (max: {self.max_file_size / (1024*1024)}MB)")
            
            if file_size == 0:
                errors.append("File is empty")
                
        except Exception as e:
            errors.append(f"Cannot access file: {e}")
        
        # Check file type support
        file_type, _ = self.detect_file_type(filepath)
        if file_type == FileType.UNKNOWN:
            errors.append("Unknown or unsupported file type")
        elif not self.supported_types.get(file_type, False):
            errors.append(f"File type {file_type.value} not supported (missing required library)")
        
        return len(errors) == 0, errors
    
    def extract_text_from_file(self, filepath: str) -> ProcessingResult:
        """Extract text from file with comprehensive error handling"""
        import time
        start_time = time.time()
        
        # Validate file
        is_valid, validation_errors = self.validate_file(filepath)
        if not is_valid:
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                text_content="",
                metadata=self._create_error_metadata(filepath),
                errors=validation_errors,
                warnings=[],
                processing_time=time.time() - start_time
            )
        
        # Detect file type
        file_type, mime_type = self.detect_file_type(filepath)
        
        # Create metadata
        metadata = self._create_metadata(filepath, file_type, mime_type)
        
        # Extract text based on file type
        try:
            if file_type == FileType.PDF:
                text, errors, warnings = self._extract_from_pdf(filepath)
            elif file_type == FileType.DOCX:
                text, errors, warnings = self._extract_from_docx(filepath)
            elif file_type == FileType.TXT:
                text, errors, warnings = self._extract_from_text(filepath)
            elif file_type == FileType.HTML:
                text, errors, warnings = self._extract_from_html(filepath)
            else:
                return ProcessingResult(
                    status=ProcessingStatus.UNSUPPORTED,
                    text_content="",
                    metadata=metadata,
                    errors=[f"File type {file_type.value} not supported"],
                    warnings=[],
                    processing_time=time.time() - start_time
                )
            
            # Update metadata with extracted content info
            metadata.word_count = len(text.split()) if text else 0
            metadata.estimated_reading_time = max(1, metadata.word_count // 200)  # ~200 WPM reading speed
            
            # Determine status
            if errors:
                status = ProcessingStatus.PARTIAL if text else ProcessingStatus.ERROR
            else:
                status = ProcessingStatus.SUCCESS
            
            return ProcessingResult(
                status=status,
                text_content=text,
                metadata=metadata,
                errors=errors,
                warnings=warnings,
                processing_time=time.time() - start_time
            )
            
        except Exception as e:
            logger.error(f"❌ Unexpected error processing {filepath}: {e}")
            return ProcessingResult(
                status=ProcessingStatus.ERROR,
                text_content="",
                metadata=metadata,
                errors=[f"Unexpected processing error: {e}"],
                warnings=[],
                processing_time=time.time() - start_time
            )
    
    def _create_metadata(self, filepath: str, file_type: FileType, mime_type: str) -> FileMetadata:
        """Create file metadata"""
        try:
            filename = os.path.basename(filepath)
            size_bytes = os.path.getsize(filepath)
            
            return FileMetadata(
                filename=filename,
                file_type=file_type,
                size_bytes=size_bytes,
                mime_type=mime_type
            )
        except Exception as e:
            logger.warning(f"⚠️ Failed to create metadata: {e}")
            return self._create_error_metadata(filepath)
    
    def _create_error_metadata(self, filepath: str) -> FileMetadata:
        """Create minimal metadata for error cases"""
        return FileMetadata(
            filename=os.path.basename(filepath) if filepath else "unknown",
            file_type=FileType.UNKNOWN,
            size_bytes=0,
            mime_type="application/octet-stream"
        )
    
    def _extract_from_pdf(self, filepath: str) -> Tuple[str, List[str], List[str]]:
        """Extract text from PDF file"""
        if not HAS_PYPDF2 or PyPDF2 is None:
            return "", ["PyPDF2 library not available"], []
        
        text_parts = []
        errors = []
        warnings = []
        
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt("")  # Try empty password
                    except:
                        return "", ["PDF is password protected"], []
                
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                        else:
                            warnings.append(f"Page {page_num + 1} appears to be empty or contains only images")
                    except Exception as e:
                        errors.append(f"Failed to extract text from page {page_num + 1}: {e}")
                
                if not text_parts and page_count > 0:
                    errors.append("No text could be extracted from PDF (may contain only images)")
                
        except Exception as e:
            return "", [f"PDF processing error: {e}"], warnings
        
        return "\n\n".join(text_parts), errors, warnings
    
    def _extract_from_docx(self, filepath: str) -> Tuple[str, List[str], List[str]]:
        """Extract text from DOCX file"""
        if not HAS_DOCX or DocxDocument is None:
            return "", ["python-docx library not available"], []
        
        text_parts = []
        errors = []
        warnings = []
        
        try:
            doc = DocxDocument(filepath)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_parts.append(f"\n[Table {table_count}]\n" + "\n".join(table_text) + "\n")
            
            if not text_parts:
                warnings.append("Document appears to be empty")
            
        except Exception as e:
            return "", [f"DOCX processing error: {e}"], warnings
        
        return "\n\n".join(text_parts), errors, warnings
    
    def _extract_from_text(self, filepath: str) -> Tuple[str, List[str], List[str]]:
        """Extract text from plain text file"""
        errors = []
        warnings = []
        
        try:
            # Try to detect encoding
            encoding = 'utf-8'
            if HAS_CHARDET and chardet is not None:
                try:
                    with open(filepath, 'rb') as f:
                        sample = f.read(8192)
                        detected = chardet.detect(sample)
                        if detected and detected.get('encoding') and detected.get('confidence', 0) > 0.7:
                            encoding = detected['encoding']
                except:
                    pass  # Fall back to utf-8
            
            # Read file with detected/default encoding
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    content = f.read()
                return content, errors, warnings
            except UnicodeDecodeError:
                # Try common encodings as fallback
                for fallback_encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(filepath, 'r', encoding=fallback_encoding) as f:
                            content = f.read()
                        warnings.append(f"Used fallback encoding: {fallback_encoding}")
                        return content, errors, warnings
                    except UnicodeDecodeError:
                        continue
                
                errors.append("Could not decode text file with any supported encoding")
                return "", errors, warnings
                
        except Exception as e:
            return "", [f"Text file processing error: {e}"], warnings
    
    def _extract_from_html(self, filepath: str) -> Tuple[str, List[str], List[str]]:
        """Extract text from HTML file (basic implementation)"""
        import re
        
        errors = []
        warnings = []
        
        try:
            # Read HTML content
            with open(filepath, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Basic HTML tag removal (not as sophisticated as BeautifulSoup)
            # Remove script and style content
            html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', ' ', html_content)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            if not text:
                warnings.append("No text content found in HTML file")
            else:
                warnings.append("Basic HTML processing used - formatting may be lost")
            
            return text, errors, warnings
            
        except Exception as e:
            return "", [f"HTML processing error: {e}"], warnings
    
    def get_file_info(self, filepath: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        try:
            file_type, mime_type = self.detect_file_type(filepath)
            
            info = {
                'filename': os.path.basename(filepath),
                'filepath': filepath,
                'file_type': file_type.value,
                'mime_type': mime_type,
                'supported': self.supported_types.get(file_type, False),
                'exists': os.path.exists(filepath)
            }
            
            if info['exists']:
                stat = os.stat(filepath)
                info.update({
                    'size_bytes': stat.st_size,
                    'size_mb': round(stat.st_size / (1024 * 1024), 2),
                    'modified_time': stat.st_mtime,
                    'is_readable': os.access(filepath, os.R_OK)
                })
                
                # Add specific info based on file type
                if file_type == FileType.PDF and HAS_PYPDF2 and PyPDF2 is not None:
                    try:
                        with open(filepath, 'rb') as f:
                            reader = PyPDF2.PdfReader(f)
                            info['page_count'] = len(reader.pages)
                            info['encrypted'] = reader.is_encrypted
                    except:
                        pass
            
            return info
            
        except Exception as e:
            logger.error(f"❌ Failed to get file info: {e}")
            return {'error': str(e)}
    
    def batch_process_files(self, filepaths: List[str]) -> Dict[str, ProcessingResult]:
        """Process multiple files and return results"""
        results = {}
        
        for filepath in filepaths:
            try:
                result = self.extract_text_from_file(filepath)
                results[filepath] = result
                logger.info(f"📄 Processed {os.path.basename(filepath)}: {result.status.value}")
            except Exception as e:
                logger.error(f"❌ Failed to process {filepath}: {e}")
                results[filepath] = ProcessingResult(
                    status=ProcessingStatus.ERROR,
                    text_content="",
                    metadata=self._create_error_metadata(filepath),
                    errors=[f"Processing failed: {e}"],
                    warnings=[],
                    processing_time=0.0
                )
        
        return results
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        extensions = []
        
        for file_type, supported in self.supported_types.items():
            if supported:
                if file_type == FileType.PDF:
                    extensions.append('.pdf')
                elif file_type == FileType.DOCX:
                    extensions.append('.docx')
                elif file_type == FileType.TXT:
                    extensions.extend(['.txt', '.text'])
                elif file_type == FileType.HTML:
                    extensions.extend(['.html', '.htm'])
        
        return extensions
    
    def create_sample_file(self, content: str, file_type: FileType, output_dir: str) -> str:
        """Create a sample file for testing purposes"""
        filename_map = {
            FileType.TXT: "sample.txt",
            FileType.HTML: "sample.html"
        }
        
        filename = filename_map.get(file_type, "sample.txt")
        filepath = os.path.join(output_dir, filename)
        
        try:
            os.makedirs(output_dir, exist_ok=True)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                if file_type == FileType.HTML:
                    f.write(f"<html><body>{content}</body></html>")
                else:
                    f.write(content)
            
            logger.info(f"✅ Created sample file: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"❌ Failed to create sample file: {e}")
            raise